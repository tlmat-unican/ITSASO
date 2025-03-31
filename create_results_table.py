import os
from docx import Document
from docx.shared import Inches

def create_table_with_title(document, folders, title, col_names, parent_folder):
    document.add_paragraph(title, style='Heading 1')
    table = document.add_table(rows=len(folders)+1, cols=len(col_names)+1)
    table.style = "Table Grid"

    # Set column headers (image file names)
    for col_idx, file_name in enumerate(col_names, start=1):
        table.cell(0, col_idx).text = file_name  # First row

     # Set row headers (folder names)
    for row_idx, folder in enumerate(folders, start=1):
        with open(os.path.join(parent_folder, "info.txt"), "r") as f:
            mapping = {line.split("\t")[0]: line.split("\t")[2].strip() for line in f.readlines()}

        table.cell(row_idx, 0).text = mapping.get(folder, folder)

        for col_idx, file_name in enumerate(col_names, start=1):
            img_path = os.path.join(parent_folder, folder, file_name)
            print(img_path)
            cell = table.cell(row_idx, col_idx)
            if os.path.exists(img_path):
                print(img_path)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(1.9685))  # 5 cm in inches
            else:
                cell.text = "Not Found"

def create_word_image_table(parent_folder, col_names, i=0, rows=20, cols=9, image_size=Inches(1.0)):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(19.685)  # 50 cm in inches
    section.page_height = Inches(19.685)  # 50 cm in inches

    static_folders = [f"sim{i}" for i in range(1, 11)]
    dynamic_folders = [f"sim{i}" for i in range(11, 21)]

    create_table_with_title(document, static_folders, "Static", col_names, parent_folder)

    document.add_paragraph()

    create_table_with_title(document, dynamic_folders, "Dynamic", col_names, parent_folder)

    document.add_paragraph()

    table2 = document.add_table(rows=3, cols=3)
    table2.style = "Table Grid"

    # Set column headers
    table2.cell(0, 1).text = "average_Battery"
    table2.cell(0, 2).text = "battery_vs_fails"

    # Set row headers
    table2.cell(1, 0).text = "static"
    table2.cell(2, 0).text = "dynamic"

    # Insert images or text into the table cells
    for row_idx in range(1, 3):
        for col_idx in range(1, 3):
            cell = table2.cell(row_idx, col_idx)
            image_name = f"{'average_Battery' if col_idx == 1 else 'battery_vs_fails_scatter'}_{row_idx-1}.png"  # Adjust image name based on row and column
            image_path = f"./sim/{image_name}"

            if os.path.exists(image_path):
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(1.9685*2))
            else:
                cell.text = "Not Found"

     # Save the document
    document.save(f"./sim/results_table_{i}.docx")


create_word_image_table("sim", ["iotDecisions.png", "DecisionsOverTime.png", "average_Battery.png", "FailedTasksTotals.png", "FailsOverTime.png", "tTotalBoxplot.png", "fogQProc.png", "Rewards_Mean.png"],i=0, rows=20, cols=8)
